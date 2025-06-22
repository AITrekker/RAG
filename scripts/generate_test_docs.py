"""
Generates sample PDF and Word documents for testing the RAG platform.
"""

import logging
from pathlib import Path
from docx import Document
from fpdf import FPDF

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Document Content ---
SAMPLE_CONTENT = {
    "title": "Project Nova Technical Specifications",
    "sections": [
        {
            "heading": "1.0 Introduction",
            "text": "Project Nova is a next-generation enterprise platform designed for high-throughput data analysis. This document outlines the core technical specifications, architecture, and performance benchmarks."
        },
        {
            "heading": "2.0 System Architecture",
            "text": "The system employs a microservices-based architecture, orchestrated using Kubernetes. The primary services include: a data ingestion pipeline, an asynchronous task queue (RabbitMQ), a metadata database (PostgreSQL), and a query processing API."
        },
        {
            "heading": "3.0 Performance",
            "text": "The data ingestion service is designed to handle up to 10,000 documents per hour. The RAG query pipeline has a target latency of less than 2 seconds for p95 queries. The embedding model is 'all-MiniLM-L6-v2' and the generative model is 'tiiuae/falcon-7b-instruct'."
        },
        {
            "heading": "4.0 Security",
            "text": "All inter-service communication is secured using mTLS. Data at rest is encrypted using AES-256. Role-Based Access Control (RBAC) is enforced at the API gateway and within each service."
        }
    ]
}

def create_word_document(output_path: Path):
    """Creates a .docx file with the sample content."""
    try:
        document = Document()
        document.add_heading(SAMPLE_CONTENT['title'], level=1)

        for section in SAMPLE_CONTENT['sections']:
            document.add_heading(section['heading'], level=2)
            document.add_paragraph(section['text'])
            document.add_paragraph() # Add a space

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        logging.info(f"Successfully created Word document: {output_path}")
    except Exception as e:
        logging.error(f"Failed to create Word document: {e}", exc_info=True)

def create_pdf_document(output_path: Path):
    """Creates a .pdf file with the sample content."""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, SAMPLE_CONTENT['title'], ln=True, align='C')
        pdf.ln(10)

        for section in SAMPLE_CONTENT['sections']:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, section['heading'], ln=True)
            pdf.set_font("Arial", '', 12)
            pdf.multi_cell(0, 5, section['text'])
            pdf.ln(5)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        pdf.output(str(output_path))
        logging.info(f"Successfully created PDF document: {output_path}")
    except Exception as e:
        logging.error(f"Failed to create PDF document: {e}", exc_info=True)

def main():
    """Main function to generate all test documents."""
    logging.info("Starting generation of test documents...")
    project_root = Path(__file__).parent.parent
    output_dir = project_root / "documents" / "default"
    
    # Define output paths
    word_output_path = output_dir / "word_docs" / "Project_Nova_Specs.docx"
    pdf_output_path = output_dir / "pdfs" / "Project_Nova_Specs.pdf"
    
    # Create the documents
    create_word_document(word_output_path)
    create_pdf_document(pdf_output_path)
    
    logging.info("Test document generation complete.")

if __name__ == "__main__":
    main() 