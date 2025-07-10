"""
LlamaIndex document processing pipeline with tenant isolation
"""

import asyncio
from typing import List, Dict, Any, Optional, Tuple
from uuid import UUID
from pathlib import Path
from dataclasses import dataclass

from sqlalchemy.ext.asyncio import AsyncSession

from src.backend.models.database import File
from src.backend.config.settings import get_settings
from .llamaindex_chunker import create_tenant_chunker, LlamaIndexChunk

settings = get_settings()


@dataclass
class ProcessedDocument:
    """Document processed through LlamaIndex pipeline"""
    file_id: UUID
    filename: str
    chunks: List[LlamaIndexChunk]
    metadata: Dict[str, Any]
    processing_stats: Dict[str, Any]


class TenantIsolatedDocumentPipeline:
    """
    LlamaIndex document processing pipeline with strict tenant isolation
    """
    
    def __init__(self, tenant_id: UUID, db_session: AsyncSession):
        self.tenant_id = tenant_id
        self.db = db_session
        self._document_loaders = {}
        self._text_extractors = {}
        self._chunker = None
        
    async def initialize(self):
        """Initialize LlamaIndex document processing components"""
        try:
            # Import LlamaIndex components
            from llama_index.readers.file import PDFReader, DocxReader, UnstructuredReader
            from llama_index.core.readers import SimpleDirectoryReader
            
            # Initialize document loaders with tenant isolation
            self._document_loaders = {
                'pdf': PDFReader(),
                'docx': DocxReader(),
                'txt': UnstructuredReader(),
                'md': UnstructuredReader()
            }
            
            # Initialize tenant-specific chunker
            self._chunker = await create_tenant_chunker(self.tenant_id)
            
            print(f"✓ LlamaIndex document pipeline initialized for tenant {self.tenant_id}")
            
        except ImportError as e:
            print(f"⚠️ LlamaIndex readers not available: {e}")
            self._document_loaders = {}
            self._chunker = None
        except Exception as e:
            print(f"⚠️ Error initializing LlamaIndex pipeline: {e}")
            self._document_loaders = {}
            self._chunker = None
    
    async def process_file(self, file_record: File) -> ProcessedDocument:
        """
        Process a single file through the LlamaIndex pipeline
        """
        file_path = f"./data/uploads/{file_record.file_path}"
        
        # Extract text using LlamaIndex or fallback
        text_content = await self._extract_text_with_llamaindex(file_path, file_record)
        
        # Process through chunker
        chunks = await self._chunker.chunk_text(
            text=text_content,
            file_id=file_record.id,
            filename=file_record.filename,
            metadata={
                'file_path': file_record.file_path,
                'mime_type': file_record.mime_type,
                'file_size': file_record.file_size,
                'tenant_id': str(self.tenant_id)
            }
        )
        
        # Generate processing statistics
        processing_stats = self._generate_processing_stats(text_content, chunks)
        
        return ProcessedDocument(
            file_id=file_record.id,
            filename=file_record.filename,
            chunks=chunks,
            metadata={
                'tenant_id': str(self.tenant_id),
                'file_path': file_record.file_path,
                'mime_type': file_record.mime_type,
                'file_size': file_record.file_size,
                'text_length': len(text_content),
                'processing_method': 'llamaindex_pipeline'
            },
            processing_stats=processing_stats
        )
    
    async def _extract_text_with_llamaindex(self, file_path: str, file_record: File) -> str:
        """Extract text using LlamaIndex readers with tenant isolation"""
        try:
            # Determine file type
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            if file_extension in self._document_loaders:
                loader = self._document_loaders[file_extension]
                
                # Load document with tenant context
                documents = loader.load_data(file=Path(file_path))
                
                # Extract text from documents
                text_content = ""
                for doc in documents:
                    # Ensure tenant isolation in document metadata
                    doc.metadata.update({
                        'tenant_id': str(self.tenant_id),
                        'file_id': str(file_record.id),
                        'filename': file_record.filename
                    })
                    
                    text_content += doc.text + "\n"
                
                print(f"✓ LlamaIndex extracted text from {file_record.filename}: {len(text_content)} chars")
                return text_content.strip()
            else:
                # Fallback to simple file reading
                return await self._fallback_text_extraction(file_path, file_record)
                
        except Exception as e:
            print(f"⚠️ LlamaIndex text extraction failed: {e}, using fallback")
            return await self._fallback_text_extraction(file_path, file_record)
    
    async def _fallback_text_extraction(self, file_path: str, file_record: File) -> str:
        """Fallback text extraction when LlamaIndex is not available"""
        try:
            # Try to read as plain text first
            if file_record.mime_type and file_record.mime_type.startswith("text/"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            # PDF extraction fallback
            if file_record.mime_type == "application/pdf":
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    return f"[PDF file: {file_record.filename} - PDF processing not available]"
            
            # DOCX extraction fallback
            if file_record.mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    return f"[DOCX file: {file_record.filename} - DOCX processing not available]"
            
            # Try as plain text anyway
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                return f"[Binary file: {file_record.filename} - text extraction not possible]"
                
        except Exception as e:
            print(f"⚠️ Fallback text extraction failed: {e}")
            return f"[Failed to extract text from: {file_record.filename}]"
    
    def _generate_processing_stats(self, text_content: str, chunks: List[LlamaIndexChunk]) -> Dict[str, Any]:
        """Generate processing statistics for the document"""
        if not chunks:
            return {
                'total_chunks': 0,
                'total_characters': 0,
                'total_words': 0,
                'total_tokens': 0,
                'avg_chunk_length': 0,
                'processing_method': 'llamaindex_pipeline'
            }
        
        total_tokens = sum(chunk.token_count for chunk in chunks)
        chunk_lengths = [len(chunk.content) for chunk in chunks]
        
        return {
            'total_chunks': len(chunks),
            'total_characters': len(text_content),
            'total_words': len(text_content.split()),
            'total_tokens': total_tokens,
            'avg_chunk_length': sum(chunk_lengths) / len(chunk_lengths),
            'min_chunk_length': min(chunk_lengths),
            'max_chunk_length': max(chunk_lengths),
            'processing_method': 'llamaindex_pipeline',
            'tenant_id': str(self.tenant_id)
        }
    
    async def process_multiple_files(self, file_records: List[File]) -> List[ProcessedDocument]:
        """
        Process multiple files concurrently with tenant isolation
        """
        if not file_records:
            return []
        
        # Create concurrent processing tasks
        tasks = []
        for file_record in file_records:
            task = self.process_file(file_record)
            tasks.append(task)
        
        # Process with controlled concurrency
        batch_size = 3  # Process 3 files at a time
        results = []
        
        for i in range(0, len(tasks), batch_size):
            batch = tasks[i:i + batch_size]
            batch_results = await asyncio.gather(*batch, return_exceptions=True)
            
            for j, result in enumerate(batch_results):
                if isinstance(result, Exception):
                    print(f"⚠️ File processing failed for {file_records[i + j].filename}: {result}")
                    # Create empty result for failed processing
                    results.append(ProcessedDocument(
                        file_id=file_records[i + j].id,
                        filename=file_records[i + j].filename,
                        chunks=[],
                        metadata={
                            'tenant_id': str(self.tenant_id),
                            'processing_error': str(result)
                        },
                        processing_stats={'error': str(result)}
                    ))
                else:
                    results.append(result)
        
        return results
    
    async def get_pipeline_statistics(self) -> Dict[str, Any]:
        """Get pipeline statistics for the tenant"""
        return {
            'tenant_id': str(self.tenant_id),
            'document_loaders_available': len(self._document_loaders),
            'supported_formats': list(self._document_loaders.keys()),
            'chunker_available': self._chunker is not None,
            'pipeline_initialized': bool(self._document_loaders and self._chunker)
        }
    
    async def validate_tenant_isolation(self, processed_docs: List[ProcessedDocument]) -> Dict[str, Any]:
        """Validate that tenant isolation is maintained across all processed documents"""
        validation_results = {
            'tenant_id': str(self.tenant_id),
            'total_documents': len(processed_docs),
            'total_chunks': sum(len(doc.chunks) for doc in processed_docs),
            'isolation_breaches': [],
            'is_valid': True
        }
        
        # Check each document and chunk for proper tenant isolation
        for doc in processed_docs:
            # Check document metadata
            doc_tenant_id = doc.metadata.get('tenant_id')
            if doc_tenant_id != str(self.tenant_id):
                validation_results['isolation_breaches'].append({
                    'type': 'document_metadata',
                    'file_id': str(doc.file_id),
                    'expected_tenant': str(self.tenant_id),
                    'actual_tenant': doc_tenant_id
                })
                validation_results['is_valid'] = False
            
            # Check chunk metadata
            for chunk in doc.chunks:
                chunk_tenant_id = chunk.metadata.get('tenant_id')
                if chunk_tenant_id != str(self.tenant_id):
                    validation_results['isolation_breaches'].append({
                        'type': 'chunk_metadata',
                        'file_id': str(doc.file_id),
                        'chunk_index': chunk.chunk_index,
                        'expected_tenant': str(self.tenant_id),
                        'actual_tenant': chunk_tenant_id
                    })
                    validation_results['is_valid'] = False
        
        return validation_results


# Factory function for creating tenant-isolated document pipelines
async def create_tenant_document_pipeline(tenant_id: UUID, db_session: AsyncSession) -> TenantIsolatedDocumentPipeline:
    """Create and initialize a tenant-isolated LlamaIndex document pipeline"""
    pipeline = TenantIsolatedDocumentPipeline(tenant_id, db_session)
    await pipeline.initialize()
    return pipeline